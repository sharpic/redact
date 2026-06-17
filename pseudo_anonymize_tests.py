#!/usr/bin/env python3
"""
pseudo_anonymize_tests.py — Unit and integration tests.

Run all:   python3 -m pytest pseudo_anonymize_tests.py -v
Run fast:  python3 -m pytest pseudo_anonymize_tests.py -v -m "not slow"
"""

import re
import sys
import tempfile
import unittest
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))

from pseudo_anonymize import (
    _apply_template,
    _overlaps,
    _passes_exclusions,
    PseudonymRegistry,
    PatternDef,
    build_patterns,
    anonymize_text,
    deanonymize_text,
    load_config,
    DEFAULT_CONFIG,
)


# ── Test fixtures ──────────────────────────────────────────────────────────────

def _pat_email():
    return PatternDef(
        name='email',
        template='person{n:03d}@anon.invalid',
        regex=re.compile(r'\b[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}\b'),
    )

def _pat_name():
    return PatternDef(
        name='name',
        template='Person{n:03d}',
        regex=re.compile(r'\b[A-Z][a-z]{1,25}(?:\s+[A-Z][a-z]{1,25}){1,3}\b'),
        exclusions=frozenset({'January', 'Monday', 'English', 'Figure', 'Table',
                              'Section', 'Page', 'True', 'False', 'Microsoft'}),
        min_word_length=2,
    )

def _pat_id():
    return PatternDef(
        name='id',
        template='preserve_length',
        regex=re.compile(r'\b\d{7,8}\b'),
    )

def _pat_username():
    return PatternDef(
        name='username',
        template='user{n:04d}xx',
        regex=re.compile(r'\b(?=[a-zA-Z0-9]*[a-zA-Z])(?=[a-zA-Z0-9]*\d)[a-zA-Z0-9]{8}\b'),
    )

def _all_patterns():
    return [_pat_email(), _pat_name(), _pat_id(), _pat_username()]

def _anon(text, patterns=None):
    """Convenience: anonymize with a fresh registry, return (result, registry)."""
    registry = PseudonymRegistry()
    result = anonymize_text(text, registry, patterns or _all_patterns(), nlp=None)
    return result, registry


# ── _apply_template ────────────────────────────────────────────────────────────

class TestApplyTemplate(unittest.TestCase):

    def test_counter_zero_padded(self):
        self.assertEqual(_apply_template('Person{n:03d}', 1,  'x'), 'Person001')
        self.assertEqual(_apply_template('Person{n:03d}', 42, 'x'), 'Person042')

    def test_counter_no_padding(self):
        self.assertEqual(_apply_template('X{n}', 7, 'x'), 'X7')

    def test_orig_substitution(self):
        self.assertEqual(_apply_template('({orig})', 1, 'secret'), '(secret)')

    def test_orig_len_substitution(self):
        self.assertEqual(_apply_template('len={orig_len}', 1, 'hello'), 'len=5')

    def test_email_template(self):
        self.assertEqual(_apply_template('person{n:03d}@anon.invalid', 5, 'x'),
                         'person005@anon.invalid')

    def test_preserve_length_7_digit(self):
        result = _apply_template('preserve_length', 1, '1234567')
        self.assertEqual(result, '0000001')
        self.assertEqual(len(result), 7)

    def test_preserve_length_8_digit(self):
        result = _apply_template('preserve_length', 99, '12345678')
        self.assertEqual(result, '00000099')
        self.assertEqual(len(result), 8)

    def test_preserve_length_large_counter(self):
        result = _apply_template('preserve_length', 10_000_000, '1234567')
        self.assertEqual(result, '10000000')   # counter > orig length — no truncation


# ── _overlaps ──────────────────────────────────────────────────────────────────

class TestOverlaps(unittest.TestCase):

    def setUp(self):
        self.spans = [(5, 10, 'A'), (20, 25, 'B')]

    def test_before_all_spans(self):
        self.assertFalse(_overlaps(self.spans, 0, 4))

    def test_between_spans(self):
        self.assertFalse(_overlaps(self.spans, 11, 19))

    def test_after_all_spans(self):
        self.assertFalse(_overlaps(self.spans, 26, 30))

    def test_adjacent_left_not_overlapping(self):
        self.assertFalse(_overlaps(self.spans, 0, 5))    # ends where span starts

    def test_adjacent_right_not_overlapping(self):
        self.assertFalse(_overlaps(self.spans, 10, 15))  # starts where span ends

    def test_overlaps_at_start(self):
        self.assertTrue(_overlaps(self.spans, 3, 7))

    def test_overlaps_at_end(self):
        self.assertTrue(_overlaps(self.spans, 8, 15))

    def test_contained_within_span(self):
        self.assertTrue(_overlaps(self.spans, 6, 9))

    def test_span_contained_within_candidate(self):
        self.assertTrue(_overlaps(self.spans, 4, 12))

    def test_empty_span_list(self):
        self.assertFalse(_overlaps([], 0, 10))


# ── _passes_exclusions ─────────────────────────────────────────────────────────

class TestPassesExclusions(unittest.TestCase):

    def setUp(self):
        self.pat = PatternDef(
            name='name', template='X',
            exclusions=frozenset({'January', 'Monday', 'Table'}),
            min_word_length=2,
        )

    def test_clean_name_passes(self):
        self.assertTrue(_passes_exclusions('John Smith', self.pat))

    def test_excluded_first_word(self):
        self.assertFalse(_passes_exclusions('January Report', self.pat))

    def test_excluded_second_word(self):
        self.assertFalse(_passes_exclusions('Sales Table', self.pat))

    def test_excluded_single_word(self):
        self.assertFalse(_passes_exclusions('Table', self.pat))

    def test_short_word_rejected_by_min_length(self):
        self.assertFalse(_passes_exclusions('Jo A', self.pat))  # 'A' is 1 char

    def test_min_length_exactly_met(self):
        self.assertTrue(_passes_exclusions('Jo Sm', self.pat))  # both ≥ 2 chars

    def test_no_exclusions_set(self):
        pat = PatternDef(name='id', template='X', regex=None)
        self.assertTrue(_passes_exclusions('anything here', pat))

    def test_min_word_length_1_allows_single_char(self):
        pat = PatternDef(name='x', template='X', min_word_length=1)
        self.assertTrue(_passes_exclusions('A B', pat))


# ── PseudonymRegistry ──────────────────────────────────────────────────────────

class TestPseudonymRegistry(unittest.TestCase):

    def setUp(self):
        self.pat = _pat_email()
        self.reg = PseudonymRegistry()

    def test_same_original_same_pseudonym(self):
        a = self.reg.get_or_create('user@x.com', self.pat)
        b = self.reg.get_or_create('user@x.com', self.pat)
        self.assertEqual(a, b)

    def test_different_originals_different_pseudonyms(self):
        a = self.reg.get_or_create('a@x.com', self.pat)
        b = self.reg.get_or_create('b@x.com', self.pat)
        self.assertNotEqual(a, b)

    def test_counter_increments_per_category(self):
        a = self.reg.get_or_create('a@x.com', self.pat)
        b = self.reg.get_or_create('b@x.com', self.pat)
        self.assertEqual(a, 'person001@anon.invalid')
        self.assertEqual(b, 'person002@anon.invalid')

    def test_whitespace_stripped_for_key(self):
        a = self.reg.get_or_create('user@x.com', self.pat)
        b = self.reg.get_or_create('  user@x.com  ', self.pat)
        self.assertEqual(a, b)

    def test_separate_counters_per_category(self):
        name_pat = _pat_name()
        self.reg.get_or_create('a@x.com', self.pat)
        self.reg.get_or_create('Alice Smith', name_pat)
        self.assertEqual(self.reg.get_or_create('b@x.com', self.pat),
                         'person002@anon.invalid')
        self.assertEqual(self.reg.get_or_create('Bob Jones', name_pat),
                         'Person002')

    def test_mapping_file_dict_inverted(self):
        self.reg.get_or_create('user@x.com', self.pat)
        d = self.reg.mapping_file_dict()
        self.assertEqual(d['person001@anon.invalid'], 'user@x.com')

    def test_stats(self):
        name_pat = _pat_name()
        self.reg.get_or_create('a@x.com', self.pat)
        self.reg.get_or_create('b@x.com', self.pat)
        self.reg.get_or_create('Alice Smith', name_pat)
        self.assertEqual(self.reg.stats(), {'email': 2, 'name': 1})


# ── anonymize_text ─────────────────────────────────────────────────────────────

class TestAnonymizeText(unittest.TestCase):

    # ── email ──────────────────────────────────────────────────────────────────

    def test_email_replaced(self):
        result, _ = _anon('Contact user@example.com today')
        self.assertNotIn('user@example.com', result)
        self.assertIn('@anon.invalid', result)

    def test_email_with_plus_addressing(self):
        result, _ = _anon('Send to user+tag@example.com please')
        self.assertNotIn('user+tag@example.com', result)
        self.assertIn('@anon.invalid', result)

    def test_email_with_subdomain(self):
        result, _ = _anon('Reply to user@mail.dept.ac.uk here')
        self.assertNotIn('@mail.dept.ac.uk', result)

    # ── IDs ────────────────────────────────────────────────────────────────────

    def test_7_digit_id_replaced(self):
        result, _ = _anon('Student 1234567 enrolled')
        self.assertNotIn('1234567', result)
        self.assertIn('0000001', result)

    def test_8_digit_id_replaced(self):
        result, _ = _anon('Ref 12345678 filed')
        self.assertNotIn('12345678', result)
        self.assertIn('00000001', result)

    def test_7_digit_preserves_length(self):
        result, _ = _anon('ID: 1234567')
        replaced = [t for t in result.split() if t.isdigit()]
        self.assertTrue(all(len(t) == 7 for t in replaced))

    def test_8_digit_preserves_length(self):
        result, _ = _anon('ID: 12345678')
        replaced = [t for t in result.split() if t.isdigit()]
        self.assertTrue(all(len(t) == 8 for t in replaced))

    def test_6_digit_not_caught(self):
        result, _ = _anon('Code 123456 here')
        self.assertIn('123456', result)

    def test_9_digit_not_caught(self):
        result, _ = _anon('Number 123456789 here')
        self.assertIn('123456789', result)

    def test_id_embedded_in_longer_number_not_caught(self):
        # \b ensures 1234567 inside 012345678 is not extracted
        result, _ = _anon('Code 012345678 here')
        self.assertIn('012345678', result)

    # ── usernames ──────────────────────────────────────────────────────────────

    def test_8_char_mixed_username_replaced(self):
        result, _ = _anon('Login: jsmith01 via portal')
        self.assertNotIn('jsmith01', result)
        self.assertIn('user0001xx', result)

    def test_all_digit_8_char_not_a_username(self):
        result, _ = _anon('Code: 12345678')
        self.assertNotIn('user', result)  # caught as id, not username

    def test_all_alpha_8_char_not_a_username(self):
        result, _ = _anon('Word: password')
        self.assertNotIn('user', result)

    def test_7_char_mixed_not_a_username(self):
        result, _ = _anon('ref: abc1234')  # only 7 chars
        self.assertNotIn('user', result)

    def test_9_char_mixed_not_a_username(self):
        result, _ = _anon('ref: abc123def')  # 9 chars
        self.assertNotIn('user', result)

    # ── names (heuristic, no spaCy) ────────────────────────────────────────────

    def test_name_heuristic_two_words(self):
        result, _ = _anon('Signed by John Smith', patterns=[_pat_name()])
        self.assertNotIn('John Smith', result)
        self.assertIn('Person', result)

    def test_name_heuristic_three_words(self):
        result, _ = _anon('Written by Mary Anne Jones', patterns=[_pat_name()])
        self.assertNotIn('Mary Anne Jones', result)

    def test_excluded_word_not_detected_as_name(self):
        result, _ = _anon('Date is January Monday')
        self.assertIn('January', result)
        self.assertIn('Monday', result)

    def test_single_title_case_word_not_caught(self):
        # A capitalised word not adjacent to another title-case word is not a name
        result, _ = _anon('Submit the Report today', patterns=[_pat_name()])
        self.assertIn('Report', result)

    # ── overlap / ordering ─────────────────────────────────────────────────────

    def test_email_not_double_detected_as_name(self):
        result, _ = _anon('Email: John.Smith@example.com here')
        self.assertEqual(result.count('@anon.invalid'), 1)
        self.assertNotIn('Person', result)

    def test_first_pattern_wins_on_overlap(self):
        # Email pattern runs before name; the email address wins
        result, _ = _anon('Contact John.Smith@example.com')
        self.assertIn('@anon.invalid', result)
        self.assertNotIn('Person', result)

    # ── consistency ────────────────────────────────────────────────────────────

    def test_consistent_replacement_across_calls(self):
        registry = PseudonymRegistry()
        pats = _all_patterns()
        r1 = anonymize_text('user@x.com', registry, pats, nlp=None)
        r2 = anonymize_text('user@x.com', registry, pats, nlp=None)
        self.assertEqual(r1, r2)

    def test_two_occurrences_get_same_pseudonym(self):
        registry = PseudonymRegistry()
        result = anonymize_text(
            'From user@x.com to user@x.com', registry, _all_patterns(), nlp=None
        )
        hits = [w for w in result.split() if '@anon.invalid' in w]
        self.assertEqual(len(hits), 2)
        self.assertEqual(hits[0], hits[1])

    def test_different_values_get_different_pseudonyms(self):
        result, _ = _anon('a@x.com and b@x.com')
        self.assertIn('person001@anon.invalid', result)
        self.assertIn('person002@anon.invalid', result)

    # ── edge cases ─────────────────────────────────────────────────────────────

    def test_empty_string(self):
        result, _ = _anon('')
        self.assertEqual(result, '')

    def test_whitespace_only(self):
        result, _ = _anon('   ')
        self.assertEqual(result, '   ')

    def test_no_pii_unchanged(self):
        text = 'No sensitive data here at all.'
        result, _ = _anon(text, patterns=[_pat_email(), _pat_id(), _pat_username()])
        self.assertEqual(result, text)

    def test_pii_at_start(self):
        result, _ = _anon('user@x.com is the contact', patterns=[_pat_email()])
        self.assertNotIn('user@x.com', result)

    def test_pii_at_end(self):
        result, _ = _anon('contact is user@x.com', patterns=[_pat_email()])
        self.assertNotIn('user@x.com', result)

    def test_multiple_pii_types_on_one_line(self):
        result, _ = _anon('user@x.com 1234567 ab12cd34')
        self.assertNotIn('user@x.com', result)
        self.assertNotIn('1234567', result)
        self.assertNotIn('ab12cd34', result)

    def test_unicode_context_untouched(self):
        result, _ = _anon('Ünïcödé context, email user@x.com here')
        self.assertIn('Ünïcödé', result)
        self.assertNotIn('user@x.com', result)


# ── deanonymize_text ───────────────────────────────────────────────────────────

class TestDeanonymizeText(unittest.TestCase):

    def setUp(self):
        self.mapping = {
            'Person001': 'John Smith',
            'person001@anon.invalid': 'john@example.com',
            '0000001': '1234567',
            'user0001xx': 'jsmith01',
        }

    def test_name_reversal(self):
        self.assertEqual(deanonymize_text('Hello Person001', self.mapping),
                         'Hello John Smith')

    def test_email_reversal(self):
        self.assertEqual(deanonymize_text('Email: person001@anon.invalid', self.mapping),
                         'Email: john@example.com')

    def test_id_reversal(self):
        self.assertEqual(deanonymize_text('ID 0000001 enrolled', self.mapping),
                         'ID 1234567 enrolled')

    def test_multiple_types_reversed(self):
        result = deanonymize_text('Person001 <person001@anon.invalid>', self.mapping)
        self.assertEqual(result, 'John Smith <john@example.com>')

    def test_empty_string(self):
        self.assertEqual(deanonymize_text('', self.mapping), '')

    def test_no_pseudonyms_unchanged(self):
        text = 'Nothing to replace here.'
        self.assertEqual(deanonymize_text(text, self.mapping), text)

    def test_longest_first_prevents_partial_match(self):
        # 'person001@anon.invalid' contains 'person001' as a substring.
        # Longest pseudonym must be replaced first or the partial match corrupts it.
        mapping = {
            'person001@anon.invalid': 'real@email.com',
            'person001': 'PARTIAL',
        }
        result = deanonymize_text('person001@anon.invalid', mapping)
        self.assertEqual(result, 'real@email.com')


# ── round-trips ────────────────────────────────────────────────────────────────

class TestRoundTrip(unittest.TestCase):
    """anonymize → deanonymize must exactly restore the original."""

    def _trip(self, text):
        registry = PseudonymRegistry()
        anon = anonymize_text(text, registry, _all_patterns(), nlp=None)
        return deanonymize_text(anon, registry.mapping_file_dict())

    def test_email(self):
        t = 'Contact user@example.com for details'
        self.assertEqual(self._trip(t), t)

    def test_7_digit_id(self):
        t = 'Student 1234567 enrolled'
        self.assertEqual(self._trip(t), t)

    def test_8_digit_id(self):
        t = 'Reference 87654321 filed'
        self.assertEqual(self._trip(t), t)

    def test_username(self):
        t = 'Login ab12cd34 active'
        self.assertEqual(self._trip(t), t)

    def test_mixed_pii(self):
        t = 'User ab12cd34 <user@dept.org> ID 1234567 enrolled'
        self.assertEqual(self._trip(t), t)

    def test_repeated_values(self):
        t = 'user@x.com and again user@x.com'
        self.assertEqual(self._trip(t), t)

    def test_no_pii_unchanged(self):
        t = 'Nothing sensitive in this sentence.'
        self.assertEqual(self._trip(t), t)

    def test_multiple_emails(self):
        t = 'From a@x.com to b@x.com cc c@x.com'
        self.assertEqual(self._trip(t), t)


# ── build_patterns ─────────────────────────────────────────────────────────────

class TestBuildPatterns(unittest.TestCase):

    def _cfg(self, patterns):
        return {'patterns': patterns}

    def test_basic_regex_compiled(self):
        pats = build_patterns(self._cfg([
            {'name': 'test', 'regex': r'\bfoo\b', 'template': 'BAR{n}'}
        ]))
        self.assertEqual(len(pats), 1)
        self.assertIsNotNone(pats[0].regex)
        self.assertIsNotNone(pats[0].regex.search('foo'))

    def test_name_and_template_set(self):
        pats = build_patterns(self._cfg([
            {'name': 'myid', 'regex': r'\d+', 'template': 'ID{n:04d}'}
        ]))
        self.assertEqual(pats[0].name, 'myid')
        self.assertEqual(pats[0].template, 'ID{n:04d}')

    def test_spacy_ner_flag(self):
        pats = build_patterns(self._cfg([
            {'name': 'name', 'spacy_ner': True, 'regex': r'\b[A-Z]\w+\b', 'template': 'X'}
        ]))
        self.assertTrue(pats[0].spacy_ner)

    def test_default_spacy_labels(self):
        pats = build_patterns(self._cfg([
            {'name': 'name', 'spacy_ner': True, 'template': 'X'}
        ]))
        self.assertEqual(pats[0].spacy_labels, ['PERSON'])

    def test_custom_spacy_labels(self):
        pats = build_patterns(self._cfg([
            {'name': 'org', 'spacy_ner': True, 'spacy_labels': ['ORG', 'GPE'], 'template': 'X'}
        ]))
        self.assertEqual(pats[0].spacy_labels, ['ORG', 'GPE'])

    def test_exclusions_become_frozenset(self):
        pats = build_patterns(self._cfg([
            {'name': 'n', 'regex': r'\w+', 'template': 'X', 'exclusions': ['January', 'Monday']}
        ]))
        self.assertIsInstance(pats[0].exclusions, frozenset)
        self.assertIn('January', pats[0].exclusions)

    def test_regex_ignorecase_flag(self):
        pats = build_patterns(self._cfg([
            {'name': 'x', 'regex': r'foo', 'template': 'X', 'regex_flags': ['IGNORECASE']}
        ]))
        self.assertTrue(pats[0].regex.flags & re.IGNORECASE)
        self.assertIsNotNone(pats[0].regex.search('FOO'))

    def test_no_regex_when_omitted(self):
        pats = build_patterns(self._cfg([
            {'name': 'name', 'spacy_ner': True, 'template': 'X'}
        ]))
        self.assertIsNone(pats[0].regex)

    def test_order_preserved(self):
        pats = build_patterns(self._cfg([
            {'name': 'a', 'regex': r'x', 'template': 'A'},
            {'name': 'b', 'regex': r'y', 'template': 'B'},
            {'name': 'c', 'regex': r'z', 'template': 'C'},
        ]))
        self.assertEqual([p.name for p in pats], ['a', 'b', 'c'])

    def test_empty_config(self):
        self.assertEqual(build_patterns({}), [])


# ── load_config ────────────────────────────────────────────────────────────────

class TestLoadConfig(unittest.TestCase):

    def test_loads_default_config(self):
        config = load_config(DEFAULT_CONFIG)
        self.assertIn('patterns', config)

    def test_default_has_email_pattern(self):
        config = load_config(DEFAULT_CONFIG)
        names = [p['name'] for p in config['patterns']]
        self.assertIn('email', names)

    def test_default_has_name_pattern(self):
        config = load_config(DEFAULT_CONFIG)
        names = [p['name'] for p in config['patterns']]
        self.assertIn('name', names)

    def test_default_has_id_pattern(self):
        config = load_config(DEFAULT_CONFIG)
        names = [p['name'] for p in config['patterns']]
        self.assertIn('id', names)

    def test_default_has_username_pattern(self):
        config = load_config(DEFAULT_CONFIG)
        names = [p['name'] for p in config['patterns']]
        self.assertIn('username', names)

    def test_name_pattern_has_spacy_ner(self):
        config = load_config(DEFAULT_CONFIG)
        name_pat = next(p for p in config['patterns'] if p['name'] == 'name')
        self.assertTrue(name_pat.get('spacy_ner', False))

    def test_id_pattern_uses_preserve_length(self):
        config = load_config(DEFAULT_CONFIG)
        id_pat = next(p for p in config['patterns'] if p['name'] == 'id')
        self.assertEqual(id_pat['template'], 'preserve_length')

    def test_missing_config_exits(self):
        with self.assertRaises(SystemExit):
            load_config(Path('/nonexistent/no_such_file.toml'))

    def test_custom_config_file(self):
        content = b'[[patterns]]\nname = "test"\nregex = \'foo\'\ntemplate = "BAR{n}"\n'
        with tempfile.NamedTemporaryFile(suffix='.toml', delete=False) as f:
            f.write(content)
            tmp = Path(f.name)
        try:
            config = load_config(tmp)
            self.assertEqual(config['patterns'][0]['name'], 'test')
        finally:
            tmp.unlink()


# ── file handler: docx ─────────────────────────────────────────────────────────

def _docx_available():
    try:
        from docx import Document  # noqa: F401
        return True
    except ImportError:
        return False


def _make_docx(path, paragraphs=(), table_rows=()):
    from docx import Document
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    if table_rows:
        table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for r, row in enumerate(table_rows):
            for c, val in enumerate(row):
                table.cell(r, c).text = val
    doc.save(str(path))


def _read_docx_para(path, index=0):
    from docx import Document
    return Document(str(path)).paragraphs[index].text


class TestFileHandlerDocx(unittest.TestCase):

    def setUp(self):
        if not _docx_available():
            self.skipTest('python-docx not installed')

    def test_paragraph_anonymized(self):
        from pseudo_anonymize import process_docx
        with tempfile.TemporaryDirectory() as d:
            src = Path(d) / 'src.docx'
            out = Path(d) / 'out.docx'
            _make_docx(src, paragraphs=['John Smith emailed user@example.com ID 1234567'])

            registry = PseudonymRegistry()
            process_docx(src, out, lambda t: anonymize_text(t, registry, _all_patterns(), nlp=None))

            text = _read_docx_para(out)
            self.assertNotIn('user@example.com', text)
            self.assertNotIn('1234567', text)

    def test_table_cells_anonymized(self):
        from docx import Document
        from pseudo_anonymize import process_docx
        with tempfile.TemporaryDirectory() as d:
            src = Path(d) / 'src.docx'
            out = Path(d) / 'out.docx'
            _make_docx(src, table_rows=[['user@example.com', '1234567']])

            registry = PseudonymRegistry()
            process_docx(src, out, lambda t: anonymize_text(t, registry, _all_patterns(), nlp=None))

            doc2 = Document(str(out))
            self.assertNotIn('user@example.com', doc2.tables[0].cell(0, 0).text)
            self.assertNotIn('1234567', doc2.tables[0].cell(0, 1).text)

    def test_round_trip(self):
        from pseudo_anonymize import process_docx
        original = 'user@example.com ID 1234567'
        with tempfile.TemporaryDirectory() as d:
            src = Path(d) / 'src.docx'
            anon = Path(d) / 'anon.docx'
            restored_path = Path(d) / 'restored.docx'
            _make_docx(src, paragraphs=[original])

            registry = PseudonymRegistry()
            process_docx(src, anon, lambda t: anonymize_text(t, registry, _all_patterns(), nlp=None))
            mapping = registry.mapping_file_dict()
            process_docx(anon, restored_path, lambda t: deanonymize_text(t, mapping))

            text = _read_docx_para(restored_path)
            self.assertIn('user@example.com', text)
            self.assertIn('1234567', text)

    def test_empty_paragraphs_untouched(self):
        from docx import Document
        from pseudo_anonymize import process_docx
        with tempfile.TemporaryDirectory() as d:
            src = Path(d) / 'src.docx'
            out = Path(d) / 'out.docx'
            _make_docx(src, paragraphs=['', 'normal text', ''])

            registry = PseudonymRegistry()
            process_docx(src, out, lambda t: anonymize_text(t, registry, _all_patterns(), nlp=None))

            paras = [p.text for p in Document(str(out)).paragraphs]
            self.assertIn('normal text', paras)


# ── file handler: xlsx ─────────────────────────────────────────────────────────

class TestFileHandlerXlsx(unittest.TestCase):

    @classmethod
    def setUpClass(cls):
        try:
            import openpyxl
            cls.openpyxl = openpyxl
        except ImportError:
            cls.openpyxl = None

    def setUp(self):
        if self.openpyxl is None:
            self.skipTest('openpyxl not installed')

    def test_string_cells_anonymized(self):
        from pseudo_anonymize import process_xlsx
        with tempfile.TemporaryDirectory() as d:
            src = Path(d) / 'src.xlsx'
            out = Path(d) / 'out.xlsx'

            wb = self.openpyxl.Workbook()
            ws = wb.active
            ws.append(['Name', 'Email', 'ID'])
            ws.append(['John Smith', 'j@example.com', '1234567'])
            wb.save(str(src))

            registry = PseudonymRegistry()
            process_xlsx(src, out, lambda t: anonymize_text(t, registry, _all_patterns(), nlp=None))

            wb2 = self.openpyxl.load_workbook(str(out))
            flat = [str(v) for row in wb2.active.iter_rows(values_only=True)
                    for v in row if v is not None]
            self.assertNotIn('j@example.com', flat)
            self.assertNotIn('1234567', flat)

    def test_numeric_cells_untouched(self):
        from pseudo_anonymize import process_xlsx
        with tempfile.TemporaryDirectory() as d:
            src = Path(d) / 'src.xlsx'
            out = Path(d) / 'out.xlsx'

            wb = self.openpyxl.Workbook()
            ws = wb.active
            ws['A1'] = 42
            ws['B1'] = 3.14
            wb.save(str(src))

            registry = PseudonymRegistry()
            process_xlsx(src, out, lambda t: anonymize_text(t, registry, _all_patterns(), nlp=None))

            wb2 = self.openpyxl.load_workbook(str(out))
            self.assertEqual(wb2.active['A1'].value, 42)
            self.assertAlmostEqual(wb2.active['B1'].value, 3.14)

    def test_multiple_sheets_processed(self):
        from pseudo_anonymize import process_xlsx
        with tempfile.TemporaryDirectory() as d:
            src = Path(d) / 'src.xlsx'
            out = Path(d) / 'out.xlsx'

            wb = self.openpyxl.Workbook()
            wb.active['A1'] = 'Sheet1: user@a.com'
            ws2 = wb.create_sheet('Sheet2')
            ws2['A1'] = 'Sheet2: user@b.com'
            wb.save(str(src))

            registry = PseudonymRegistry()
            process_xlsx(src, out, lambda t: anonymize_text(t, registry, _all_patterns(), nlp=None))

            wb2 = self.openpyxl.load_workbook(str(out))
            self.assertNotIn('user@a.com', str(wb2.active['A1'].value))
            self.assertNotIn('user@b.com', str(wb2['Sheet2']['A1'].value))

    def test_round_trip(self):
        from pseudo_anonymize import process_xlsx
        with tempfile.TemporaryDirectory() as d:
            src = Path(d) / 'src.xlsx'
            anon = Path(d) / 'anon.xlsx'
            rest = Path(d) / 'rest.xlsx'

            wb = self.openpyxl.Workbook()
            ws = wb.active
            ws.append(['j@example.com', '1234567', 'ab12cd34'])
            wb.save(str(src))

            registry = PseudonymRegistry()
            process_xlsx(src, anon, lambda t: anonymize_text(t, registry, _all_patterns(), nlp=None))
            mapping = registry.mapping_file_dict()
            process_xlsx(anon, rest, lambda t: deanonymize_text(t, mapping))

            wb3 = self.openpyxl.load_workbook(str(rest))
            flat = [str(v) for row in wb3.active.iter_rows(values_only=True) for v in row if v]
            self.assertIn('j@example.com', flat)
            self.assertIn('1234567', flat)
            self.assertIn('ab12cd34', flat)


# ── file handler: odt ──────────────────────────────────────────────────────────

def _odt_available():
    try:
        from odf.opendocument import OpenDocumentText  # noqa: F401
        return True
    except ImportError:
        return False


def _make_odt(path, text):
    from odf.opendocument import OpenDocumentText
    from odf.text import P
    doc = OpenDocumentText()
    doc.text.addElement(P(text=text))
    doc.save(str(path))


def _read_odt(path):
    from odf.opendocument import load as odf_load
    doc = odf_load(str(path))
    parts = []
    def _walk(el):
        for child in el.childNodes:
            if hasattr(child, 'data'):
                parts.append(child.data)
            else:
                _walk(child)
    _walk(doc.text)
    return ' '.join(parts)


class TestFileHandlerOdt(unittest.TestCase):

    def setUp(self):
        if not _odt_available():
            self.skipTest('odfpy not installed')

    def test_text_anonymized(self):
        from pseudo_anonymize import process_odt
        with tempfile.TemporaryDirectory() as d:
            src = Path(d) / 'src.odt'
            out = Path(d) / 'out.odt'
            _make_odt(src, 'Contact user@example.com ID 1234567')

            registry = PseudonymRegistry()
            process_odt(src, out, lambda t: anonymize_text(t, registry, _all_patterns(), nlp=None))

            content = _read_odt(out)
            self.assertNotIn('user@example.com', content)
            self.assertNotIn('1234567', content)

    def test_round_trip(self):
        from pseudo_anonymize import process_odt
        with tempfile.TemporaryDirectory() as d:
            src = Path(d) / 'src.odt'
            anon = Path(d) / 'anon.odt'
            rest = Path(d) / 'rest.odt'
            _make_odt(src, 'user@example.com 1234567')

            registry = PseudonymRegistry()
            process_odt(src, anon, lambda t: anonymize_text(t, registry, _all_patterns(), nlp=None))
            mapping = registry.mapping_file_dict()
            process_odt(anon, rest, lambda t: deanonymize_text(t, mapping))

            content = _read_odt(rest)
            self.assertIn('user@example.com', content)
            self.assertIn('1234567', content)


# ── file handler: pdf ──────────────────────────────────────────────────────────

class TestFileHandlerPdf(unittest.TestCase):

    def setUp(self):
        try:
            import pdfplumber  # noqa: F401
        except ImportError:
            self.skipTest('pdfplumber not installed')
        # Use the PDF that lives next to the script; skip if absent
        self.pdf = Path(__file__).parent / 'Academic Calendar 2026-27.pdf'
        if not self.pdf.exists():
            self.skipTest('no test PDF found in project directory')

    def test_produces_text_file(self):
        from pseudo_anonymize import process_pdf
        with tempfile.TemporaryDirectory() as d:
            out = Path(d) / 'out.txt'
            registry = PseudonymRegistry()
            process_pdf(self.pdf, out,
                        lambda t: anonymize_text(t, registry, _all_patterns(), nlp=None))
            self.assertTrue(out.exists())
            self.assertGreater(out.stat().st_size, 0)

    def test_emails_replaced_in_output(self):
        from pseudo_anonymize import process_pdf
        with tempfile.TemporaryDirectory() as d:
            out = Path(d) / 'out.txt'
            registry = PseudonymRegistry()
            process_pdf(self.pdf, out,
                        lambda t: anonymize_text(t, registry, _all_patterns(), nlp=None))
            text = out.read_text(encoding='utf-8')
            # The calendar contains manchester.ac.uk addresses
            self.assertNotIn('@manchester.ac.uk', text)
            self.assertIn('@anon.invalid', text)


# ── spaCy NER (skipped unless model is present) ────────────────────────────────

try:
    import spacy as _spacy
    _spacy.load('en_core_web_sm')
    _SPACY_OK = True
except Exception:
    _SPACY_OK = False


@unittest.skipUnless(_SPACY_OK, 'spaCy en_core_web_sm not installed')
class TestSpacyNER(unittest.TestCase):

    @classmethod
    def setUpClass(cls):
        import spacy
        cls.nlp = spacy.load('en_core_web_sm')
        cls.patterns = _all_patterns()

    def _run(self, text):
        registry = PseudonymRegistry()
        return anonymize_text(text, registry, self.patterns, nlp=self.nlp), registry

    def test_person_name_detected(self):
        result, _ = self._run('The report was written by Alice Johnson.')
        self.assertNotIn('Alice Johnson', result)
        self.assertIn('Person', result)

    def test_email_still_replaced(self):
        result, _ = self._run('Email alice@example.com for details.')
        self.assertNotIn('alice@example.com', result)
        self.assertIn('@anon.invalid', result)

    def test_name_and_email_both_replaced(self):
        result, _ = self._run('Alice Johnson <alice@example.com> submitted')
        self.assertNotIn('Alice Johnson', result)
        self.assertNotIn('alice@example.com', result)

    def test_date_not_replaced(self):
        # Dates are tagged DATE by spaCy, not PERSON — should be unchanged
        result, _ = self._run('The deadline is 15 March 2025')
        self.assertNotIn('Person', result)
        self.assertIn('2025', result)

    def test_round_trip_with_spacy(self):
        original = 'Alice Johnson <alice@example.com> ID 1234567'
        registry = PseudonymRegistry()
        anon = anonymize_text(original, registry, self.patterns, nlp=self.nlp)
        restored = deanonymize_text(anon, registry.mapping_file_dict())
        self.assertEqual(restored, original)

    def test_id_and_username_still_caught(self):
        result, _ = self._run('ID 1234567 login ab12cd34')
        self.assertNotIn('1234567', result)
        self.assertNotIn('ab12cd34', result)


# ── entry point ────────────────────────────────────────────────────────────────

if __name__ == '__main__':
    unittest.main(verbosity=2)
