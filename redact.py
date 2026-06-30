#!/usr/bin/env python3
"""
redact.py — Redact PII in documents.

Replaces personally identifiable information (PII) in a document with
consistent placeholders, and writes a JSON mapping file that can later
be used to restore the originals.

Supported input formats
-----------------------
  .docx   Microsoft Word
  .odt    LibreOffice / OpenDocument Text
  .xlsx   Microsoft Excel
  .pdf    PDF (output is plain text — formatting cannot be preserved)

Detected PII (default config)
------------------------------
  email      user@domain.tld                → person001@anon.invalid
  name       Alice Smith                    → Person001
  id         1234567  /  12345678           → same-width zero-padded number
  username   ab12cd34  (8-char alphanumeric)→ user0001xx

All patterns are defined in redact_config.toml.
Add a [[patterns]] block there to extend detection — no Python changes needed.

Output files
------------
  <stem>_redacted.<ext>   redacted copy of the document
  <stem>_mapping.json     pseudonym → original  (keep this safe; it reverses the process)

Install dependencies
--------------------
  pip install python-docx odfpy openpyxl pdfplumber
  pip install spacy && python -m spacy download en_core_web_sm  # recommended for names
"""

import re
import json
import sys
import tomllib
import argparse
from dataclasses import dataclass, field
from pathlib import Path
from datetime import datetime


# ── Config loading ─────────────────────────────────────────────────────────────

DEFAULT_CONFIG = Path(__file__).with_name('redact_config.toml')


def load_config(path: Path) -> dict:
    """Load and return the TOML config dict, exiting with an error if not found."""
    if not path.exists():
        sys.exit(
            f'Error: config file not found: {path}\n'
            f'       Expected redact_config.toml next to the script,\n'
            f'       or specify one with --config <file>.'
        )
    with open(path, 'rb') as f:
        return tomllib.load(f)


@dataclass
class PatternDef:
    name: str
    template: str
    regex: re.Pattern | None = None
    spacy_ner: bool = False
    spacy_labels: list = field(default_factory=lambda: ['PERSON'])
    exclusions: frozenset = frozenset()
    min_word_length: int = 1
    pn_only: bool = False


def build_patterns(config: dict) -> list[PatternDef]:
    """Compile each [[patterns]] entry from the config into a PatternDef.

    Regex strings are compiled here so misconfigured patterns fail at startup,
    not mid-document.  Optional regex_flags are OR-ed in from the re module.
    """
    patterns = []
    for p in config.get('patterns', []):
        flags = 0
        for flag_name in p.get('regex_flags', []):
            flags |= getattr(re, flag_name)

        compiled = re.compile(p['regex'], flags) if 'regex' in p else None

        patterns.append(PatternDef(
            name=p['name'],
            template=p['template'],
            regex=compiled,
            spacy_ner=p.get('spacy_ner', False),
            spacy_labels=p.get('spacy_labels', ['PERSON']),
            exclusions=frozenset(p.get('exclusions', [])),
            min_word_length=p.get('min_word_length', 1),
            pn_only=p.get('pn_only', False),
        ))
    return patterns


# ── PseudonymRegistry ──────────────────────────────────────────────────────────

def _apply_template(template: str, n: int, original: str) -> str:
    """Render a pseudonym from a template string and a sequential counter.

    Supports {n} (counter), {orig} (original text), {orig_len} (its length),
    and the special keyword 'preserve_length' which zero-pads n to match the
    digit count of the original value.
    """
    if template == 'preserve_length':
        return str(n).zfill(len(original))
    return template.format(n=n, orig=original, orig_len=len(original))


class PseudonymRegistry:
    """Consistent original→pseudonym mapping; never re-uses a pseudonym."""

    def __init__(self):
        self._map = {}       # (pattern_name, original_stripped) → pseudonym
        self._counters = {}

    def get_or_create(self, original: str, pat: PatternDef) -> str:
        key = (pat.name, original.strip())
        if key not in self._map:
            n = self._counters.get(pat.name, 0) + 1
            self._counters[pat.name] = n
            self._map[key] = _apply_template(pat.template, n, original)
        return self._map[key]

    def mapping_file_dict(self) -> dict:
        """Return {pseudonym: original} — the restoration key."""
        return {v: k[1] for k, v in self._map.items()}

    def stats(self) -> dict:
        return dict(self._counters)


# ── Core text redaction ────────────────────────────────────────────────────────

def _overlaps(spans: list, start: int, end: int) -> bool:
    """Return True if (start, end) overlaps any span already in the list."""
    return any(s < end and e > start for s, e, _ in spans)


def _passes_exclusions(text: str, pat: PatternDef) -> bool:
    """Return True if none of the words in text are in the pattern's exclusion set
    and every word meets the minimum length requirement."""
    words = text.split()
    if pat.min_word_length > 1 and any(len(w) < pat.min_word_length for w in words):
        return False
    return not any(w in pat.exclusions for w in words)


def redact_text(text: str, registry: PseudonymRegistry,
                patterns: list[PatternDef], nlp=None) -> str:
    """Scan text for PII using each pattern in order and replace with pseudonyms.

    Patterns are applied in config order; the first pattern to claim a span
    prevents later patterns from overlapping it (so email beats name for
    'John.Smith@example.com').  spaCy NER is used when nlp is provided and
    the pattern has spacy_ner=true; otherwise the regex fallback is used.
    Exclusions are applied to both NER and regex matches.
    """
    if not text or not text.strip():
        return text

    spans = []  # (start, end, replacement_string)

    def _add(start, end, original, pat):
        if not _overlaps(spans, start, end):
            spans.append((start, end, registry.get_or_create(original, pat)))

    for pat in patterns:
        if pat.spacy_ner and nlp is not None:
            doc = nlp(text)
            for ent in doc.ents:
                if ent.label_ in pat.spacy_labels and _passes_exclusions(ent.text, pat):
                    _add(ent.start_char, ent.end_char, ent.text, pat)
        elif pat.regex is not None:
            for m in pat.regex.finditer(text):
                if _passes_exclusions(m.group(), pat):
                    _add(m.start(), m.end(), m.group(), pat)

    if not spans:
        return text

    # Apply in reverse order so earlier offsets stay valid
    spans.sort(key=lambda x: x[0], reverse=True)
    chars = list(text)
    for start, end, repl in spans:
        chars[start:end] = list(repl)
    return ''.join(chars)


# ── Core text restoration ──────────────────────────────────────────────────────

def restore_text(text: str, pseudo_to_orig: dict) -> str:
    """Replace every pseudonym in text with its original value.

    Replacements are applied longest-first so that 'person001@anon.invalid'
    is replaced before 'person001', preventing partial corruption.
    """
    if not text:
        return text
    for pseudo, orig in sorted(pseudo_to_orig.items(), key=lambda x: -len(x[0])):
        text = text.replace(pseudo, orig)
    return text


# ── File handlers ──────────────────────────────────────────────────────────────

def _process_docx_para(para, transform_fn):
    """Redact one paragraph by concatenating all runs, transforming, then
    writing the result into the first run and clearing the rest.

    Note: this preserves paragraph-level formatting (style, alignment) but
    collapses per-run character formatting (bold/italic on individual words)
    within any replaced span.
    """
    if not para.runs:
        return
    full = ''.join(r.text for r in para.runs)
    transformed = transform_fn(full)
    if transformed != full:
        para.runs[0].text = transformed
        for r in para.runs[1:]:
            r.text = ''


def process_docx(input_path: Path, output_path: Path, transform_fn):
    """Apply transform_fn to every text span in a .docx file and save the result."""
    from docx import Document
    doc = Document(str(input_path))

    def walk(container):
        for para in container.paragraphs:
            _process_docx_para(para, transform_fn)
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    walk(cell)

    walk(doc)
    for section in doc.sections:
        for hf in (section.header, section.footer):
            if hf:
                walk(hf)

    doc.save(str(output_path))


def process_odt(input_path: Path, output_path: Path, transform_fn):
    """Apply transform_fn to every text node in a .odt file and save the result."""
    from odf.opendocument import load as odf_load

    doc = odf_load(str(input_path))

    def walk(element):
        for child in list(element.childNodes):
            if hasattr(child, 'data'):
                transformed = transform_fn(child.data)
                if transformed != child.data:
                    child.data = transformed
            else:
                walk(child)

    walk(doc.text)
    doc.save(str(output_path))


def _build_component_words(registry: 'PseudonymRegistry', patterns: list) -> set:
    """Return the set of single words that unambiguously identify one registered name.

    Used by process_xlsx (-mcn) to detect name fragments in individual cells so they
    get their own pseudonym even when the pattern regex or NER wouldn't catch them alone.
    A word is included only if it appears in exactly one registered multi-word name —
    words shared across multiple names (common first names, etc.) are excluded.
    Each detected fragment gets a fresh pseudonym via get_or_create, so de-redaction
    maps every cell back to its own original value with no collisions.
    """
    pat_names = {p.name for p in patterns if p.spacy_ner or p.name in ('name', 'proper_noun')}
    entries = [
        orig
        for (pname, orig) in registry._map
        if pname in pat_names and ' ' in orig
    ]
    component_words: set = set()
    for orig in entries:
        for word in orig.split():
            if len(word) >= 2:
                word_re = re.compile(r'(?<!\w)' + re.escape(word) + r'(?!\w)')
                if sum(1 for o in entries if word_re.search(o)) == 1:
                    component_words.add(word)
    return component_words


_FIRST_NAME_HEADERS = frozenset({
    'first name', 'firstname', 'forename', 'given name', 'christian name', 'first',
})
_LAST_NAME_HEADERS = frozenset({
    'last name', 'lastname', 'surname', 'family name', 'second name', 'last',
})
_MIDDLE_NAME_HEADERS = frozenset({
    'middle name', 'middle', 'middle initial',
})


def _find_name_column_triples(headers: list) -> list[tuple]:
    """Return (first_idx, middle_idx_or_None, last_idx) for detected name column groups."""
    norm = [str(h).lower().strip() if h else '' for h in headers]
    firsts  = [i for i, h in enumerate(norm) if h in _FIRST_NAME_HEADERS]
    lasts   = [i for i, h in enumerate(norm) if h in _LAST_NAME_HEADERS]
    middles = [i for i, h in enumerate(norm) if h in _MIDDLE_NAME_HEADERS]
    triples = []
    for fi, li in zip(firsts, lasts):
        mid = next((m for m in middles if fi < m < li), None)
        triples.append((fi, mid, li))
    return triples


def _xlsx_prescan_sheet(rows, prescan_fn):
    """Pre-scan all data rows so NER sees full names before any cell is redacted."""
    name_triples = _find_name_column_triples([cell.value for cell in rows[0]])
    for row in rows[1:]:
        str_vals = [c.value for c in row if isinstance(c.value, str) and c.value.strip()]
        if str_vals:
            prescan_fn(' '.join(str_vals))
        for fi, mi, li in name_triples:
            indices = [fi, mi, li] if mi is not None else [fi, li]
            parts = [row[i].value for i in indices
                     if i < len(row) and isinstance(row[i].value, str) and row[i].value.strip()]
            if len(parts) >= 2:
                prescan_fn(' '.join(parts))


def _xlsx_component_info(registry, patterns):
    """Return (component_words, name_pat) for the post-prescan fallback step."""
    if not registry or not patterns:
        return set(), None
    return _build_component_words(registry, patterns), next(
        (p for p in patterns if p.name == 'name'), None
    )


def _xlsx_transform_cell(cell, transform_fn, registry, name_pat, component_words):
    """Apply transform_fn to one string cell, with component-word fallback."""
    if not isinstance(cell.value, str):
        return
    result = transform_fn(cell.value)
    if result == cell.value and name_pat and cell.value.strip() in component_words:
        result = registry.get_or_create(cell.value.strip(), name_pat)
    if result != cell.value:
        cell.value = result


def process_xlsx(input_path: Path, output_path: Path, transform_fn,
                 prescan_fn=None, registry=None, patterns=None):
    """Apply transform_fn to every string cell across all sheets and save the result.
    Numeric and formula cells are left untouched.

    When prescan_fn + registry + patterns are provided (-mcn mode):
      1. Row-combination pre-scan: join all string cells per row so NER sees full names.
      2. Header-aware pre-scan: combine First Name / Last Name column pairs explicitly.
      3. Component words: after pre-scanning, identify unambiguous name fragments and
         assign each its own fresh pseudonym via get_or_create — so every cell has a
         unique mapping entry and de-redaction restores each cell to its own original.
    """
    import openpyxl
    wb = openpyxl.load_workbook(str(input_path))

    for ws in wb.worksheets:
        rows = list(ws.iter_rows())
        if not rows:
            continue

        component_words, name_pat = set(), None
        if prescan_fn is not None and len(rows) > 1:
            _xlsx_prescan_sheet(rows, prescan_fn)
            component_words, name_pat = _xlsx_component_info(registry, patterns)

        for row in rows:
            for cell in row:
                _xlsx_transform_cell(cell, transform_fn, registry, name_pat, component_words)

    wb.save(str(output_path))


def process_pdf(input_path: Path, output_path: Path, transform_fn):
    """Extract text from each PDF page, apply transform_fn line-by-line, and
    write the result as plain text.  PDF formatting is not preserved.
    """
    import pdfplumber

    pages = []
    with pdfplumber.open(str(input_path)) as pdf:
        for i, page in enumerate(pdf.pages, 1):
            raw = page.extract_text() or ''
            lines = [transform_fn(line) for line in raw.splitlines()]
            pages.append(f'--- Page {i} ---\n' + '\n'.join(lines))

    output_path.write_text('\n\n'.join(pages), encoding='utf-8')
    print(f'[note] PDF output is plain text (formatting cannot be preserved): {output_path}')


_HANDLERS = {
    '.docx': (process_docx, '.docx', 'python-docx'),
    '.odt':  (process_odt,  '.odt',  'odfpy'),
    '.xlsx': (process_xlsx, '.xlsx', 'openpyxl'),
    '.pdf':  (process_pdf,  '.txt',  'pdfplumber'),
}

_RESTORE_SUFFIX = {'.docx': '.docx', '.odt': '.odt', '.xlsx': '.xlsx', '.txt': '.txt'}


# ── Commands ───────────────────────────────────────────────────────────────────

def _resolve_redact_paths(args, out_suffix):
    """Return (output_path, mapping_path) from args, applying default naming."""
    input_path = Path(args.input)
    output_path = (Path(args.output) if args.output
                   else input_path.with_name(input_path.stem + '_redacted' + out_suffix))
    mapping_path = (Path(args.mapping) if args.mapping
                    else input_path.with_name(input_path.stem + '_mapping.json'))
    return output_path, mapping_path


def _load_nlp(config, patterns, no_spacy):
    """Load spaCy model if any pattern needs NER; return nlp or None."""
    if not any(p.spacy_ner for p in patterns) or no_spacy:
        return None
    model = config.get('settings', {}).get('spacy_model', 'en_core_web_sm')
    try:
        import spacy
        nlp = spacy.load(model)
        print(f'[info] spaCy NER active ({model}).')
        return nlp
    except ImportError:
        print('[info] spaCy not installed — using regex heuristic for names.')
        print('       pip install spacy && python -m spacy download', model)
    except OSError:
        print(f'[info] spaCy model "{model}" not found — using regex heuristic.')
        print('       python -m spacy download', model)
    return None


def _save_mapping(mapping_path, input_path, output_path, registry, nlp):
    """Write the pseudonym mapping JSON file."""
    mapping_doc = {
        'source_file': str(input_path.resolve()),
        'redacted_file': str(output_path.resolve()),
        'created': datetime.now().isoformat(timespec='seconds'),
        'spacy_ner_used': nlp is not None,
        'stats': registry.stats(),
        'mapping': registry.mapping_file_dict(),
    }
    mapping_path.write_text(
        json.dumps(mapping_doc, indent=2, ensure_ascii=False), encoding='utf-8'
    )


def _print_redact_stats(registry, nlp):
    """Print a summary of what was replaced."""
    stats = registry.stats()
    total = sum(stats.values())
    if total:
        print(f'Replaced:    {total} items — ' +
              ', '.join(f'{v} {k}(s)' for k, v in stats.items()))
    else:
        print('[note] No PII detected.')
        if not nlp:
            print('       Install spaCy for better name detection.')


def cmd_redact(args):
    input_path = Path(args.input)
    if not input_path.exists():
        sys.exit(f'Error: not found: {input_path}')

    suffix = input_path.suffix.lower()
    if suffix not in _HANDLERS:
        sys.exit(f"Error: unsupported type '{suffix}'. Supported: {', '.join(_HANDLERS)}")

    handler, out_suffix, lib = _HANDLERS[suffix]
    output_path, mapping_path = _resolve_redact_paths(args, out_suffix)

    config = load_config(Path(args.config))
    patterns = build_patterns(config)
    if not patterns:
        sys.exit('Error: no [[patterns]] defined in config file.')
    if not args.proper_nouns:
        patterns = [p for p in patterns if not p.pn_only]

    nlp = _load_nlp(config, patterns, args.no_spacy)
    registry = PseudonymRegistry()
    transform = lambda text: redact_text(text, registry, patterns, nlp)

    print(f'Redacting:   {input_path}')
    try:
        if suffix == '.xlsx' and args.multi_col_names:
            prescan = lambda text: redact_text(text, registry, patterns, nlp)
            process_xlsx(input_path, output_path, transform,
                         prescan_fn=prescan, registry=registry, patterns=patterns)
        else:
            handler(input_path, output_path, transform)
    except ImportError as exc:
        sys.exit(f'Missing library — run: pip install {lib}\n  ({exc})')

    _save_mapping(mapping_path, input_path, output_path, registry, nlp)
    print(f'Output:      {output_path}')
    print(f'Mapping:     {mapping_path}')
    _print_redact_stats(registry, nlp)


def cmd_restore(args):
    input_path = Path(args.input)
    if not input_path.exists():
        sys.exit(f'Error: not found: {input_path}')

    mapping_path = Path(args.mapping) if args.mapping else None
    if mapping_path is None:
        stem = input_path.stem
        if stem.endswith('_redacted'):
            stem = stem[:-9]
        mapping_path = input_path.with_name(stem + '_mapping.json')

    if not mapping_path.exists():
        sys.exit(f'Error: mapping file not found: {mapping_path}\n'
                 f'       Specify it with --mapping <file>')

    mapping_doc = json.loads(mapping_path.read_text(encoding='utf-8'))
    pseudo_to_orig = mapping_doc.get('mapping', {})
    if not pseudo_to_orig:
        sys.exit('Error: mapping file contains no entries.')

    suffix = input_path.suffix.lower()
    if suffix not in _HANDLERS and suffix != '.txt':
        sys.exit(f"Error: unsupported type '{suffix}'.")

    out_suffix = _RESTORE_SUFFIX.get(suffix, suffix)
    output_path = (Path(args.output) if args.output
                   else input_path.with_name(input_path.stem + '_restored' + out_suffix))

    transform = lambda text: restore_text(text, pseudo_to_orig)

    print(f'Restoring:   {input_path}')
    try:
        if suffix == '.txt':
            output_path.write_text(
                transform(input_path.read_text(encoding='utf-8')), encoding='utf-8'
            )
        else:
            handler, _, lib = _HANDLERS[suffix]
            handler(input_path, output_path, transform)
    except ImportError as exc:
        _, _, lib = _HANDLERS.get(suffix, (None, None, 'unknown'))
        sys.exit(f'Missing library — run: pip install {lib}\n  ({exc})')

    print(f'Restored:    {output_path}')
    print(f'Applied {len(pseudo_to_orig)} mapping(s) from {mapping_path}')


# ── Entry point ────────────────────────────────────────────────────────────────

_EXAMPLES = """
examples:
  # Redact a Word document (output: report_redacted.docx + report_mapping.json)
  python3 redact.py report.docx

  # Redact an Excel spreadsheet with custom output paths
  python3 redact.py data.xlsx -o data_safe.xlsx -m data_keys.json

  # Redact a PDF (output is plain text — PDF formatting cannot be preserved)
  python3 redact.py notes.pdf

  # Restore — auto-locates the mapping file from the filename
  python3 redact.py report_redacted.docx --restore
  # → report_redacted_restored.docx

  # Restore with an explicit mapping file
  python3 redact.py report_redacted.docx -r -m /secure/report_mapping.json

  # Use a custom pattern config instead of the default
  python3 redact.py report.docx --config my_patterns.toml

  # Also redact person names / proper nouns
  python3 redact.py report.docx --proper-nouns
  python3 redact.py report.docx -pn          # short form
  python3 redact.py report.docx -rn          # alias (real names)

  # Excel: link names split across columns (First Name + Last Name → same pseudonym)
  python3 redact.py data.xlsx --multi-col-names
  python3 redact.py data.xlsx -mcn           # short form

  # Disable spaCy NER when using --proper-nouns (uses regex heuristic instead)
  python3 redact.py report.docx -pn --no-spacy

detected PII  (default — edit redact_config.toml to extend):
  email      user@domain.tld              → person001@anon.invalid
  name       Alice Smith / Alice Johnson  → Person001  (spaCy NER or 2-word heuristic)
  id         1234567  /  12345678         → 0000001  /  00000001  (same digit count)
  username   ab12cd34  (8-char mixed)     → user0001xx

with -pn / --proper-nouns (adds broad proper-noun detection):
  org/place  Google / Microsoft / Oxford  → Entity001  (spaCy ORG/GPE or Title Case word)

output files:
  <stem>_redacted.<ext>   redacted copy of the document
  <stem>_mapping.json     pseudonym→original map — keep this safe and separate
                          from the redacted file; it is the only way to reverse
                          the process

adding new patterns:
  Append a [[patterns]] block to redact_config.toml.  No Python
  changes are needed.  Example — UK National Insurance numbers:

    [[patterns]]
    name    = "ni_number"
    regex   = '\\b[A-CEGHJ-PR-TW-Z]{2}\\d{6}[A-D]\\b'
    regex_flags = ["IGNORECASE"]
    template = "NI{n:04d}"
"""


def main():
    parser = argparse.ArgumentParser(
        description=(
            'Redact PII in Word (.docx), ODT, Excel (.xlsx), and PDF files.\n'
            'Patterns are configured in redact_config.toml — add a [[patterns]]\n'
            'block there to extend detection without changing any Python code.'
        ),
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=_EXAMPLES,
    )
    parser.add_argument('input', help='Input file (.docx / .odt / .xlsx / .pdf)')
    parser.add_argument('--output', '-o', metavar='FILE',
                        help='Output path (default: <input>_redacted.<ext>)')
    parser.add_argument('--mapping', '-m', metavar='FILE',
                        help='Mapping JSON path (default: <input>_mapping.json)')
    parser.add_argument('--config', '-c', metavar='FILE',
                        default=str(DEFAULT_CONFIG),
                        help=f'Config TOML path (default: {DEFAULT_CONFIG.name} next to script)')
    parser.add_argument('--proper-nouns', '--real-names', '-pn', '-rn',
                        dest='proper_nouns', action='store_true',
                        help='Also redact proper nouns and person names (uses spaCy NER when available)')
    parser.add_argument('--multi-col-names', '-mcn', dest='multi_col_names',
                        action='store_true',
                        help='Excel only: pre-scan rows to link names split across columns '
                             '(e.g. "Alice" in one cell, "Johnson" in another → same pseudonym). '
                             'Uses both row-combination and header-aware detection '
                             '(First Name / Last Name column headers).')
    parser.add_argument('--no-spacy', action='store_true',
                        help='Disable spaCy NER when using --proper-nouns (uses regex heuristic instead)')
    parser.add_argument('--restore', '-r', action='store_true',
                        help='Reverse: restore originals from mapping file')

    args = parser.parse_args()

    if args.restore:
        cmd_restore(args)
    else:
        cmd_redact(args)


if __name__ == '__main__':  # pragma: no cover
    main()
